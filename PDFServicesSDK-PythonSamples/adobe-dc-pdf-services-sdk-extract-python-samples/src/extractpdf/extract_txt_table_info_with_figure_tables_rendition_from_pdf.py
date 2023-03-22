# Copyright 2021 Adobe. All rights reserved.
# This file is licensed to you under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License. You may obtain a copy
# of the License at http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software distributed under
# the License is distributed on an "AS IS" BASIS, WITHOUT WARRANTIES OR REPRESENTATIONS
# OF ANY KIND, either express or implied. See the License for the specific language
# governing permissions and limitations under the License.

import logging
import os.path
import zipfile
import json
from json2html import *
import docx
import re
import fnmatch

zip_file = "C:/Users/anand/Documents/python/pdfextract/PDFServicesSDK-PythonSamples/adobe-dc-pdf-services-sdk-extract-python-samples/output/ExtractTextTableWithFigureTableRendition.zip"
destination = 'C:/Users/anand/Documents/python/pdfextract/PDFServicesSDK-PythonSamples/adobe-dc-pdf-services-sdk-extract-python-samples/output/ExtractTextTableWithFigureTableRendition/'


from adobe.pdfservices.operation.auth.credentials import Credentials
from adobe.pdfservices.operation.exception.exceptions import ServiceApiException, ServiceUsageException, SdkException
from adobe.pdfservices.operation.pdfops.options.extractpdf.extract_pdf_options import ExtractPDFOptions
from adobe.pdfservices.operation.pdfops.options.extractpdf.extract_renditions_element_type import \
    ExtractRenditionsElementType
from adobe.pdfservices.operation.pdfops.options.extractpdf.extract_element_type import ExtractElementType
from adobe.pdfservices.operation.execution_context import ExecutionContext
from adobe.pdfservices.operation.io.file_ref import FileRef
from adobe.pdfservices.operation.pdfops.extract_pdf_operation import ExtractPDFOperation


logging.basicConfig(level=os.environ.get("LOGLEVEL", "INFO"))

try:
    # get base path.
    base_path = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    # Initial setup, create credentials instance.
    credentials = Credentials.service_account_credentials_builder() \
        .from_file(base_path + "/pdfservices-api-credentials.json") \
        .build()

    # Create an ExecutionContext using credentials and create a new operation instance.
    execution_context = ExecutionContext.create(credentials)
    extract_pdf_operation = ExtractPDFOperation.create_new()

    # Set operation input from a source file.
    source = FileRef.create_from_local_file(base_path + "/resources/Prot_000.pdf")
    extract_pdf_operation.set_input(source)

    # Build ExtractPDF options and set them into the operation
    extract_pdf_options: ExtractPDFOptions = ExtractPDFOptions.builder() \
        .with_elements_to_extract([ExtractElementType.TEXT, ExtractElementType.TABLES]) \
        .with_elements_to_extract_renditions([ExtractRenditionsElementType.TABLES,
                                              ExtractRenditionsElementType.FIGURES]) \
        .build()
    extract_pdf_operation.set_options(extract_pdf_options)

    # Execute the operation.
    result: FileRef = extract_pdf_operation.execute(execution_context)

    # Save the result to the specified location.
    result.save_as(base_path + "/output/ExtractTextTableWithFigureTableRendition.zip")


    extract = zipfile.ZipFile(zip_file, 'r')
    extract.extractall(destination)
    extract.close()
    print("test")
    archive = zipfile.ZipFile(zip_file, 'r')
    jsonentry = archive.open('structuredData.json')
    jsondata = jsonentry.read()
    data = json.loads(jsondata)
    #formatted_table = json2html.convert(json = data)
    #index= open("index.html","w")
    #index.write(formatted_table)
    #index.close()
    cont =''
    img =''


    mydoc = docx.Document()
    #for element in data["elements"]:
        #mydoc.add_paragraph(element["Path"])
    for element in data["elements"]:

        #if(element["Path"].endswith("Table") or (re.search("Table\[.\]",element["Path"]) != None)):
            #img=''.join(re.search(".xlsx",element["filePaths"])
            #path = base_path + "/output/ExtractTextTableWithFigureTableRendition/" + img
            #mydoc.add_picture(path)
        if(element["Path"].endswith("Figure") or (re.search("Figure\[.+\]",element["Path"]) != None)):
            img_tab=''.join(element["filePaths"])
            path_tab = base_path + "/output/ExtractTextTableWithFigureTableRendition/" + img_tab
            mydoc.add_picture(path_tab)
        if(element["Path"].endswith("Title") or element["Path"].endswith("H1")  or element["Path"].endswith("P") or element["Path"].endswith("LBody") or element["Path"].endswith("Lb1") or  element["Path"].endswith("H2") or  element["Path"].endswith("Reference") or  element["Path"].endswith("Sub") or   (re.search("Sub\[.+\]",element["Path"]) != None) or  (re.search("P\[.+\]",element["Path"]) != None) or element["Path"].endswith("ParagraphSpan") or (re.search("ParagraphSpan\[.+\]",element["Path"]) != None) or (re.search("H1\[.+\]",element["Path"]) != None)  or (re.search("H2\[.+\]",element["Path"]) != None) or (re.search("Span\[.+\]",element["Path"]) != None) or element["Path"].endswith("Span")):
            cont = element["Text"]
            mydoc.add_paragraph(cont)
    mydoc.save("C:/Users/anand/Documents/python/pdfextract/PDFServicesSDK-PythonSamples/adobe-dc-pdf-services-sdk-extract-python-samples/output/my_written_file.docx")
except (ServiceApiException, ServiceUsageException, SdkException):
    logging.exception("Exception encountered while executing operation")
